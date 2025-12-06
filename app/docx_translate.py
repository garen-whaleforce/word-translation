"""
DOCX translation module using OpenAI API.
Reads a Word document, translates content to Traditional Chinese,
and writes the result to a new document.
"""

import asyncio
import os
import re
from dataclasses import dataclass, field
from enum import Enum
from typing import Callable, Optional, Awaitable, Union, Tuple

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from openai import AsyncOpenAI, AsyncAzureOpenAI, APIError, RateLimitError

# Type alias for the client (can be either OpenAI or Azure OpenAI)
OpenAIClient = Union[AsyncOpenAI, AsyncAzureOpenAI]


class TranslationError(Exception):
    """Raised when translation fails."""
    pass


class ElementType(Enum):
    """Type of document element."""
    PARAGRAPH = "paragraph"
    TABLE_CELL = "table_cell"
    HEADER = "header"
    FOOTER = "footer"
    TEXT_BOX = "text_box"


@dataclass
class TextElement:
    """Represents a text element in the document with its location."""
    element_type: ElementType
    paragraph_index: Optional[int] = None
    table_index: Optional[int] = None
    row_index: Optional[int] = None
    cell_index: Optional[int] = None
    cell_paragraph_index: Optional[int] = None
    section_index: Optional[int] = None  # For headers/footers
    header_type: Optional[str] = None  # 'default', 'first', 'even'
    textbox_index: Optional[int] = None  # For text boxes
    textbox_para_index: Optional[int] = None  # Paragraph index within text box
    xml_element: Optional[object] = None  # Store XML element reference for text boxes
    original_text: str = ""
    translated_text: str = ""


@dataclass
class TranslationStats:
    """Statistics for translation process."""
    original_chars: int = 0
    translated_chars: int = 0
    prompt_tokens: int = 0
    completion_tokens: int = 0
    total_tokens: int = 0
    api_calls: int = 0

    @property
    def estimated_cost(self) -> float:
        """Estimate cost based on GPT-4o-mini pricing (as of 2025)."""
        # GPT-4o-mini: $0.15/1M input, $0.60/1M output
        input_cost = (self.prompt_tokens / 1_000_000) * 0.15
        output_cost = (self.completion_tokens / 1_000_000) * 0.60
        return input_cost + output_cost


# System prompt for translation
SYSTEM_PROMPT = """You are a senior bilingual technical translator. Your ONLY task is to translate from **English to Traditional Chinese (Taiwan)**.

The documents are CB / IEC safety test reports and power electronics specifications. Your translation MUST sound like it was written by an experienced compliance engineer familiar with IEC/EN standards and safety reports used in Taiwan.

### Core rules
1. **Direction:** Always translate **from English to Traditional Chinese**. Never translate Chinese back to English.
2. **Style:**
   - Use formal, concise wording suitable for test reports, specifications, and certification documents.
   - Use clear engineering wording, not marketing language.
   - Keep sentence structure close to the source when it improves traceability in audits or cross-checking.
3. **Formatting & layout:**
   - Preserve tables, item numbers, headings, clause numbers, units, symbols, and values.
   - Do NOT change numbers, limits, dates, test results, verdicts, or standard identifiers.
   - Keep IEC / EN / UL standard codes (e.g., "IEC 62368-1") in English.
   - If the input contains multiple paragraphs separated by special markers like "|||", preserve these markers in your output.

4. **What must remain in English:**
   - Standard names and numbers (IEC/EN/UL/CSA, etc.).
   - Trade names, model names, company names, PCB designators (R1, C2, T1, etc.).
   - Keep abbreviations like "CB", "ICT", "AV" if they are part of standard terminology in the report.

5. **Do NOT leave English untranslated**
   - Except for items listed above, **everything else must be translated into Traditional Chinese**.
   - If you must keep a term in English for technical accuracy, add a clear Traditional Chinese explanation on first occurrence.

### Terminology – MANDATORY glossary (English ➜ Traditional Chinese)
When these English terms or phrases appear, you MUST use EXACTLY the following translations.
Always match the **longest phrase first** (e.g., match "primary winding" before the single word "primary").

Parts / components:
- Bleeding resistor ➜ 洩放電阻
- Electrolytic capacitor ➜ 電解電容
- MOSFET ➜ 電晶體
- Current limit resistor ➜ 限流電阻
- Varistor / MOV ➜ 突波吸收器
- Primary wire ➜ 一次側引線
- Line choke / Line chock ➜ 電感
- Bobbin ➜ 線架
- Plug holder ➜ 刃片插座塑膠材質
- AC connector ➜ AC 連接器
- Fuse ➜ 保險絲
- Triple insulated wire ➜ 三層絕緣線
- Trace (PCB) ➜ 銅箔

Circuit sides & windings:
- primary winding ➜ 一次側繞線
- primary circuit ➜ 一次側電路
- primary (alone, referring to primary side) ➜ 一次側
- secondary ➜ 二次側
- Sec. (abbreviation) ➜ 二次側
- winding (general) ➜ 繞線
- core (magnetic core) ➜ 鐵芯

Test conditions, environment, status:
- Unit shutdown immediately ➜ 設備立即中斷
- Unit shutdown ➜ 設備中斷
- Ambient (temperature, condition) ➜ 室溫
- Plastic enclosure outside near ➜ 塑膠外殼內側靠近
- For model ➜ 適用型號
- Optional ➜ 可選
- Interchangeable ➜ 不限
- Minimum / at least ➜ 至少

Additional wording constraints:
- NEVER translate "primary" as "初級" or "一次測" or "一次"; always use **一次側**.
- NEVER translate "secondary" as "次級"; always use **二次側**.
- Use **Traditional Chinese** characters only.

### Table cell formatting rules
- Flammability rating cells: When you see "UL 94, UL 746C" or similar, output ONLY "UL 94" (remove UL 746C)
- Empty or blank cells: Keep them empty/blank, do NOT add any content
- Certification/approval cells with file numbers: Remove file numbers, keep ONLY the certification standard names
  Example: "VDE↓40029550↓UL E249609" → "VDE" (remove all file numbers like 40029550, E249609, E121562, etc.)

### Quality checks
Before finalizing each answer, mentally check:
1. All English technical content (except standard names, model names, etc.) has been translated into Traditional Chinese.
2. All glossary terms above are applied consistently, prioritizing the longest phrase match.
3. Numbers, units, limits, clause numbers, table structures, and verdicts are preserved exactly.
4. The overall tone is that of a professional safety/compliance report used in Taiwan.

Output ONLY the translated Traditional Chinese text (with the preserved structure), without explanations."""


CHUNK_SIZE = 1500  # characters per chunk
MAX_RETRIES = 3
RETRY_DELAY = 2  # seconds
MAX_CONCURRENT_CHUNKS = 20  # max concurrent API calls per file
MAX_CONCURRENT_FILES = 2  # max concurrent file translations

# Global semaphore for file-level concurrency control
_file_semaphore: Optional[asyncio.Semaphore] = None


def get_file_semaphore() -> asyncio.Semaphore:
    """Get or create the global file semaphore."""
    global _file_semaphore
    if _file_semaphore is None:
        _file_semaphore = asyncio.Semaphore(MAX_CONCURRENT_FILES)
    return _file_semaphore


# Special translations for test result indicators
SPECIAL_TRANSLATIONS = {
    "P": "符合",
    "N/A": "不適用",
    "--": "--",
}


def get_special_translation(text: str) -> Optional[str]:
    """
    Check if text matches a special translation case.

    Args:
        text: The text to check.

    Returns:
        The special translation if matched, None otherwise.
    """
    stripped = text.strip()
    return SPECIAL_TRANSLATIONS.get(stripped)


# Patterns that indicate LLM refusal/error messages
REFUSAL_PATTERNS = [
    "抱歉，我無法",
    "抱歉，我不能",
    "對不起，我無法",
    "對不起，我不能",
    "很抱歉，我無法",
    "很抱歉，我不能",
    "I'm sorry",
    "I cannot",
    "I can't",
]


def filter_refusal_message(text: str) -> str:
    """
    Filter LLM refusal messages and replace with '--'.

    Args:
        text: The translated text to check.

    Returns:
        '--' if text is a refusal message, otherwise the original text.
    """
    if not text:
        return text

    stripped = text.strip()
    for pattern in REFUSAL_PATTERNS:
        if stripped.startswith(pattern):
            return "--"

    return text


# Type alias for progress callback: (stage: str, current: int, total: int) -> Awaitable[None]
ProgressCallback = Callable[[str, int, int], Awaitable[None]]


def get_openai_client() -> Tuple[OpenAIClient, str]:
    """
    Create and return an OpenAI client (Azure or standard).

    Returns:
        Tuple of (client, model_or_deployment_name)

    Raises:
        TranslationError: If required API keys are not set.
    """
    # Check for Azure OpenAI first
    azure_api_key = os.environ.get("AZURE_OPENAI_API_KEY")
    azure_endpoint = os.environ.get("AZURE_OPENAI_ENDPOINT")
    azure_deployment = os.environ.get("AZURE_OPENAI_DEPLOYMENT")
    azure_api_version = os.environ.get("AZURE_OPENAI_API_VERSION", "2024-12-01-preview")

    if azure_api_key and azure_endpoint and azure_deployment:
        client = AsyncAzureOpenAI(
            api_key=azure_api_key,
            azure_endpoint=azure_endpoint,
            api_version=azure_api_version
        )
        return client, azure_deployment

    # Fallback to standard OpenAI
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise TranslationError(
            "No OpenAI API key configured. "
            "Set either AZURE_OPENAI_* or OPENAI_API_KEY environment variables."
        )
    return AsyncOpenAI(api_key=api_key), "gpt-4o-mini"


async def translate_text_chunk(
    client: OpenAIClient,
    model: str,
    text: str,
    stats: TranslationStats
) -> str:
    """
    Translate a chunk of text using OpenAI API.

    Args:
        client: The OpenAI async client (Azure or standard).
        model: The model or deployment name to use.
        text: The text to translate.
        stats: TranslationStats object to update.

    Returns:
        The translated text in Traditional Chinese.

    Raises:
        TranslationError: If translation fails after retries.
    """
    if not text.strip():
        return text

    last_error = None
    for attempt in range(MAX_RETRIES):
        try:
            response = await client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": text}
                ],
                max_completion_tokens=4096
            )

            # Update stats
            stats.api_calls += 1
            if response.usage:
                stats.prompt_tokens += response.usage.prompt_tokens
                stats.completion_tokens += response.usage.completion_tokens
                stats.total_tokens += response.usage.total_tokens

            # Get content and filter refusal messages
            content = response.choices[0].message.content or ""
            return filter_refusal_message(content)

        except RateLimitError as e:
            last_error = e
            if attempt < MAX_RETRIES - 1:
                await asyncio.sleep(RETRY_DELAY * (attempt + 1))
            continue

        except APIError as e:
            if e.status_code and e.status_code >= 500:
                last_error = e
                if attempt < MAX_RETRIES - 1:
                    await asyncio.sleep(RETRY_DELAY * (attempt + 1))
                continue
            raise TranslationError(f"OpenAI API error: {str(e)}")

        except Exception as e:
            raise TranslationError(f"Unexpected error during translation: {str(e)}")

    raise TranslationError(
        f"Translation failed after {MAX_RETRIES} attempts: {str(last_error)}"
    )


def _get_all_tables(doc: Document) -> list[Table]:
    """
    Recursively get all tables including nested tables from the document.
    """
    tables = []

    def add_tables_from_element(element):
        """Recursively find tables in an element."""
        if hasattr(element, 'tables'):
            for table in element.tables:
                tables.append(table)
                # Check for nested tables in cells
                for row in table.rows:
                    for cell in row.cells:
                        add_tables_from_element(cell)

    add_tables_from_element(doc)
    return tables


def extract_text_elements(doc: Document) -> list[TextElement]:
    """
    Extract all text elements from a Word document.

    Args:
        doc: The python-docx Document object.

    Returns:
        A list of TextElement objects with location and text information.
    """
    elements: list[TextElement] = []
    processed_cells = set()  # Track processed cells to avoid duplicates

    # Extract paragraphs (not in tables)
    for para_idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:
            elements.append(TextElement(
                element_type=ElementType.PARAGRAPH,
                paragraph_index=para_idx,
                original_text=text
            ))

    # Extract table cells - get all tables including nested ones
    all_tables = _get_all_tables(doc)

    for table_idx, table in enumerate(all_tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                # Skip duplicate cells (merged cells return same object)
                cell_id = id(cell)
                if cell_id in processed_cells:
                    continue
                processed_cells.add(cell_id)

                for para_idx, paragraph in enumerate(cell.paragraphs):
                    text = paragraph.text.strip()
                    if text:
                        elements.append(TextElement(
                            element_type=ElementType.TABLE_CELL,
                            table_index=table_idx,
                            row_index=row_idx,
                            cell_index=cell_idx,
                            cell_paragraph_index=para_idx,
                            original_text=text
                        ))

    # Extract headers and footers from all sections
    for section_idx, section in enumerate(doc.sections):
        # Header types: default, first_page, even_page
        header_types = [
            ('default', section.header),
            ('first', section.first_page_header),
            ('even', section.even_page_header),
        ]
        for header_type, header in header_types:
            if header and header.is_linked_to_previous is False:
                for para_idx, paragraph in enumerate(header.paragraphs):
                    text = paragraph.text.strip()
                    if text:
                        elements.append(TextElement(
                            element_type=ElementType.HEADER,
                            section_index=section_idx,
                            header_type=header_type,
                            paragraph_index=para_idx,
                            original_text=text
                        ))

        # Footer types: default, first_page, even_page
        footer_types = [
            ('default', section.footer),
            ('first', section.first_page_footer),
            ('even', section.even_page_footer),
        ]
        for footer_type, footer in footer_types:
            if footer and footer.is_linked_to_previous is False:
                for para_idx, paragraph in enumerate(footer.paragraphs):
                    text = paragraph.text.strip()
                    if text:
                        elements.append(TextElement(
                            element_type=ElementType.FOOTER,
                            section_index=section_idx,
                            header_type=footer_type,
                            paragraph_index=para_idx,
                            original_text=text
                        ))

    # Extract text from text boxes (w:txbxContent elements)
    # These are often used in PDF conversions for floating text
    try:
        # Use lxml etree for xpath with namespaces
        nsmap = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        }

        # Find all text box content elements in the document body
        body_xml = doc.element.body
        textbox_contents = body_xml.findall('.//' + qn('w:txbxContent'))

        for tb_idx, txbx_content in enumerate(textbox_contents):
            # Find all paragraph elements within the text box
            para_elements = txbx_content.findall('.//' + qn('w:p'))

            for para_idx, para_elem in enumerate(para_elements):
                # Get all text from this paragraph
                text_elements = para_elem.findall('.//' + qn('w:t'))
                text_parts = [t.text for t in text_elements if t.text]
                text = ''.join(text_parts).strip()

                if text:
                    elements.append(TextElement(
                        element_type=ElementType.TEXT_BOX,
                        textbox_index=tb_idx,
                        textbox_para_index=para_idx,
                        xml_element=para_elem,
                        original_text=text
                    ))
    except Exception as e:
        # If text box extraction fails, continue without it
        pass

    return elements


def create_chunks(elements: list[TextElement]) -> list[list[int]]:
    """
    Group text elements into chunks for batch translation.

    Args:
        elements: List of TextElement objects.

    Returns:
        List of lists, where each inner list contains indices of elements in that chunk.
    """
    chunks: list[list[int]] = []
    current_chunk: list[int] = []
    current_length = 0

    for idx, element in enumerate(elements):
        text_length = len(element.original_text)

        if current_length + text_length > CHUNK_SIZE and current_chunk:
            chunks.append(current_chunk)
            current_chunk = []
            current_length = 0

        current_chunk.append(idx)
        current_length += text_length

    if current_chunk:
        chunks.append(current_chunk)

    return chunks


async def translate_elements(
    client: OpenAIClient,
    model: str,
    elements: list[TextElement],
    chunks: list[list[int]],
    stats: TranslationStats,
    progress_callback: Optional[ProgressCallback] = None
) -> None:
    """
    Translate all text elements in chunks with concurrent processing.

    Args:
        client: The OpenAI async client (Azure or standard).
        model: The model or deployment name to use.
        elements: List of TextElement objects to translate.
        chunks: List of index groups for batch translation.
        stats: TranslationStats object to update.
        progress_callback: Optional callback for progress updates.
    """
    # First pass: handle special translations (P, N/A, --, etc.)
    for element in elements:
        special = get_special_translation(element.original_text)
        if special is not None:
            element.translated_text = special

    separator = " ||| "
    total_chunks = len(chunks)
    completed_count = 0
    count_lock = asyncio.Lock()
    semaphore = asyncio.Semaphore(MAX_CONCURRENT_CHUNKS)

    async def process_chunk(chunk_indices: list[int]) -> None:
        """Process a single chunk with semaphore control."""
        nonlocal completed_count

        async with semaphore:
            # Filter out elements that already have special translations
            indices_to_translate = [
                idx for idx in chunk_indices
                if not elements[idx].translated_text
            ]

            if not indices_to_translate:
                # Update progress even for skipped chunks
                async with count_lock:
                    completed_count += 1
                    if progress_callback:
                        await progress_callback("translating", completed_count, total_chunks)
                return

            # Combine texts in this chunk
            texts = [elements[i].original_text for i in indices_to_translate]
            combined_text = separator.join(texts)

            # Translate the combined text
            translated_combined = await translate_text_chunk(client, model, combined_text, stats)

            # Split the translation back
            translated_texts = translated_combined.split(separator)

            # Handle case where separator might be translated or modified
            if len(translated_texts) != len(indices_to_translate):
                # Fallback: translate each element individually
                for idx in indices_to_translate:
                    translated = await translate_text_chunk(
                        client, model, elements[idx].original_text, stats
                    )
                    elements[idx].translated_text = translated
            else:
                # Assign translations to elements
                for i, idx in enumerate(indices_to_translate):
                    elements[idx].translated_text = translated_texts[i].strip()

            # Update progress
            async with count_lock:
                completed_count += 1
                if progress_callback:
                    await progress_callback("translating", completed_count, total_chunks)

    # Process all chunks concurrently
    await asyncio.gather(*[process_chunk(chunk) for chunk in chunks])


# Pattern to match sequences of dots (3 or more dots, with optional spaces and colons)
DOTS_PATTERN = re.compile(r'[.\u2026]{3,}[\s:：]*')

# Pattern to detect significant English text (ignoring standard codes, numbers, etc.)
# This matches English words that are likely to need translation
ENGLISH_WORD_PATTERN = re.compile(r'[A-Za-z]{4,}')

# Words/patterns to exclude from English detection (standards, abbreviations, etc.)
ENGLISH_EXCLUSIONS = {
    'iec', 'en', 'ul', 'csa', 'vde', 'tuv', 'cb', 'ict', 'mosfet', 'pcb',
    'ac', 'dc', 'led', 'usb', 'hdmi', 'wifi', 'http', 'https', 'api',
    'pass', 'fail', 'n/a', 'max', 'min', 'typ', 'nom', 'ref', 'see',
    'table', 'figure', 'note', 'page', 'item', 'model', 'type', 'class',
}


def _contains_significant_english(text: str) -> bool:
    """
    Check if text contains significant English words that should be translated.
    Returns True if the text has English words that are likely untranslated content.
    """
    # Find all English words (4+ letters)
    words = ENGLISH_WORD_PATTERN.findall(text.lower())

    # Filter out excluded words
    significant_words = [w for w in words if w not in ENGLISH_EXCLUSIONS]

    # If there are significant English words, text likely needs translation
    return len(significant_words) > 0


def find_untranslated_elements(doc: Document) -> list[TextElement]:
    """
    Scan the document to find elements with remaining English text.
    Returns a list of TextElement objects that need translation.
    """
    elements: list[TextElement] = []
    processed_cells = set()

    # Check paragraphs
    for para_idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text and _contains_significant_english(text):
            elements.append(TextElement(
                element_type=ElementType.PARAGRAPH,
                paragraph_index=para_idx,
                original_text=text
            ))

    # Check table cells
    all_tables = _get_all_tables(doc)
    for table_idx, table in enumerate(all_tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_id = id(cell)
                if cell_id in processed_cells:
                    continue
                processed_cells.add(cell_id)

                for para_idx, paragraph in enumerate(cell.paragraphs):
                    text = paragraph.text.strip()
                    if text and _contains_significant_english(text):
                        elements.append(TextElement(
                            element_type=ElementType.TABLE_CELL,
                            table_index=table_idx,
                            row_index=row_idx,
                            cell_index=cell_idx,
                            cell_paragraph_index=para_idx,
                            original_text=text
                        ))

    # Check headers and footers
    for section_idx, section in enumerate(doc.sections):
        for header_type, header in [('default', section.header), ('first', section.first_page_header), ('even', section.even_page_header)]:
            if header and header.is_linked_to_previous is False:
                for para_idx, paragraph in enumerate(header.paragraphs):
                    text = paragraph.text.strip()
                    if text and _contains_significant_english(text):
                        elements.append(TextElement(
                            element_type=ElementType.HEADER,
                            section_index=section_idx,
                            header_type=header_type,
                            paragraph_index=para_idx,
                            original_text=text
                        ))

        for footer_type, footer in [('default', section.footer), ('first', section.first_page_footer), ('even', section.even_page_footer)]:
            if footer and footer.is_linked_to_previous is False:
                for para_idx, paragraph in enumerate(footer.paragraphs):
                    text = paragraph.text.strip()
                    if text and _contains_significant_english(text):
                        elements.append(TextElement(
                            element_type=ElementType.FOOTER,
                            section_index=section_idx,
                            header_type=footer_type,
                            paragraph_index=para_idx,
                            original_text=text
                        ))

    # Check text boxes
    try:
        body_xml = doc.element.body
        textbox_contents = body_xml.findall('.//' + qn('w:txbxContent'))
        for tb_idx, txbx_content in enumerate(textbox_contents):
            para_elements = txbx_content.findall('.//' + qn('w:p'))
            for para_idx, para_elem in enumerate(para_elements):
                text_elements = para_elem.findall('.//' + qn('w:t'))
                text_parts = [t.text for t in text_elements if t.text]
                text = ''.join(text_parts).strip()
                if text and _contains_significant_english(text):
                    elements.append(TextElement(
                        element_type=ElementType.TEXT_BOX,
                        textbox_index=tb_idx,
                        textbox_para_index=para_idx,
                        xml_element=para_elem,
                        original_text=text
                    ))
    except Exception:
        pass

    return elements


def _clean_text(text: str) -> str:
    """Clean up text by removing dot sequences and extra whitespace."""
    # Remove dot sequences (like "......" or "…………")
    cleaned = DOTS_PATTERN.sub('', text)
    # Clean up extra whitespace
    cleaned = ' '.join(cleaned.split())
    return cleaned.strip()


def _update_paragraph_text(paragraph, text: str) -> None:
    """Helper to update paragraph text while preserving formatting."""
    # Clean the text before writing
    cleaned_text = _clean_text(text)
    if paragraph.runs:
        first_run = paragraph.runs[0]
        for run in paragraph.runs[1:]:
            run.text = ""
        first_run.text = cleaned_text
    else:
        paragraph.text = cleaned_text


def write_translations_to_doc(doc: Document, elements: list[TextElement]) -> None:
    """
    Write translated text back to the document.

    Args:
        doc: The python-docx Document object.
        elements: List of TextElement objects with translations.
    """
    # Get all tables (same order as extraction)
    all_tables = _get_all_tables(doc)

    for element in elements:
        if not element.translated_text:
            continue

        if element.element_type == ElementType.PARAGRAPH:
            if element.paragraph_index is not None:
                paragraph = doc.paragraphs[element.paragraph_index]
                _update_paragraph_text(paragraph, element.translated_text)

        elif element.element_type == ElementType.TABLE_CELL:
            if (element.table_index is not None and
                element.row_index is not None and
                element.cell_index is not None and
                element.cell_paragraph_index is not None):

                table = all_tables[element.table_index]
                cell = table.rows[element.row_index].cells[element.cell_index]
                paragraph = cell.paragraphs[element.cell_paragraph_index]
                _update_paragraph_text(paragraph, element.translated_text)

        elif element.element_type == ElementType.HEADER:
            if (element.section_index is not None and
                element.header_type is not None and
                element.paragraph_index is not None):

                section = doc.sections[element.section_index]
                if element.header_type == 'default':
                    header = section.header
                elif element.header_type == 'first':
                    header = section.first_page_header
                else:
                    header = section.even_page_header

                if header and element.paragraph_index < len(header.paragraphs):
                    paragraph = header.paragraphs[element.paragraph_index]
                    _update_paragraph_text(paragraph, element.translated_text)

        elif element.element_type == ElementType.FOOTER:
            if (element.section_index is not None and
                element.header_type is not None and
                element.paragraph_index is not None):

                section = doc.sections[element.section_index]
                if element.header_type == 'default':
                    footer = section.footer
                elif element.header_type == 'first':
                    footer = section.first_page_footer
                else:
                    footer = section.even_page_footer

                if footer and element.paragraph_index < len(footer.paragraphs):
                    paragraph = footer.paragraphs[element.paragraph_index]
                    _update_paragraph_text(paragraph, element.translated_text)

        elif element.element_type == ElementType.TEXT_BOX:
            # Update text box content directly via XML
            if element.xml_element is not None:
                cleaned_text = _clean_text(element.translated_text)

                # Find all text elements in this paragraph and update them
                text_elements = element.xml_element.findall('.//' + qn('w:t'))
                if text_elements:
                    # Clear all text elements except the first one
                    for t_elem in text_elements[1:]:
                        t_elem.text = ""
                    # Set the first text element to the translated text
                    text_elements[0].text = cleaned_text


async def translate_docx_to_zh_hant(
    src_docx_path: str,
    dst_docx_path: str,
    progress_callback: Optional[ProgressCallback] = None,
    first_pass_path: Optional[str] = None
) -> TranslationStats:
    """
    Read a DOCX file, translate its content to Traditional Chinese,
    and save the result to a new file.

    Uses file-level semaphore to limit concurrent file translations.
    Performs a second pass to catch any remaining untranslated English text.

    Args:
        src_docx_path: Path to the source DOCX file.
        dst_docx_path: Path where the translated DOCX will be saved.
        progress_callback: Optional callback for progress updates.
        first_pass_path: Optional path to save the first pass result for debugging.

    Returns:
        TranslationStats with usage statistics.

    Raises:
        TranslationError: If translation fails.
    """
    stats = TranslationStats()
    file_semaphore = get_file_semaphore()

    async with file_semaphore:
        try:
            # Report: extracting text
            if progress_callback:
                await progress_callback("extracting", 0, 0)

            # Load the document
            doc = Document(src_docx_path)

            # Extract all text elements
            elements = extract_text_elements(doc)

            if not elements:
                # No text to translate, just save a copy
                doc.save(dst_docx_path)
                return stats

            # Calculate original character count
            stats.original_chars = sum(len(e.original_text) for e in elements)

            # Create chunks for batch translation
            chunks = create_chunks(elements)

            # Get OpenAI client and translate
            client, model = get_openai_client()
            await translate_elements(client, model, elements, chunks, stats, progress_callback)

            # Calculate translated character count
            stats.translated_chars = sum(len(e.translated_text) for e in elements)

            # Report: saving document
            if progress_callback:
                await progress_callback("saving", 0, 0)

            # Write translations back to document
            write_translations_to_doc(doc, elements)

            # Save first pass result for debugging if path provided
            if first_pass_path:
                doc.save(first_pass_path)

            # === SECOND PASS: Find and translate remaining English text ===
            if progress_callback:
                await progress_callback("second_pass", 0, 0)

            # Find elements with remaining English text
            untranslated = find_untranslated_elements(doc)

            if untranslated:
                # Create chunks for second pass
                second_chunks = create_chunks(untranslated)

                # Translate remaining elements
                if progress_callback:
                    await progress_callback("translating_pass2", 0, len(second_chunks))

                await translate_elements(client, model, untranslated, second_chunks, stats, progress_callback)

                # Write second pass translations
                write_translations_to_doc(doc, untranslated)

                # Update stats
                stats.translated_chars += sum(len(e.translated_text) for e in untranslated if e.translated_text)

            # Save the final translated document
            doc.save(dst_docx_path)

            return stats

        except TranslationError:
            raise

        except Exception as e:
            raise TranslationError(f"Failed to process document: {str(e)}")
